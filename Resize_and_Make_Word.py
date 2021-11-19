#======== Import Library =====================================
from PIL import Image
import os, shutil
from docx import Document
from docx.shared import Inches, Mm, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL

def main():
    #======== Input path of all project ==========================
    print("Input Path of Project. (ex. D:_Work/_LSS/01_UAV/03_Project/โคก หนอง นา/_Round_06/)")
    input_path = input("Path : ")
    print("")

    #======== Global Variable ====================================
    photo_path = input_path +"\\Photo\\"
    os.mkdir(input_path + "\\Buffer\\")
    os.mkdir(input_path + "\\Output\\")
    buffer_Path = input_path + "\\Buffer\\"
    output_path = input_path + "\\Output\\ภาพประกอบจุดรังวัด แปลง "

    #======== Resize Photo ======================================
    dirss = os.listdir(photo_path)
    print("\nResizing...\n")

    for item_Name in dirss:
        os.mkdir(buffer_Path + item_Name)
        fieldcode = item_Name + "\\"
        a = photo_path + item_Name + "\\"
        dirs = os.listdir(a)
        print("\t"+a)
        for item in dirs:
            if os.path.isfile(a + item) and (item != 'Thumbs.db'):
                im = Image.open(a + item)
                width, height = im.size
                f, e = os.path.splitext(item)
                if(width > height):
                    imResize = im.resize((886,665), Image.ANTIALIAS)
                else:
                    imResize = im.resize((665,886), Image.ANTIALIAS)
                # imResize = im.resize((int(width/4),int(height/4)), Image.ANTIALIAS)
                imResize.save(buffer_Path + fieldcode + f + '.jpg', 'JPEG', quality=100)

    #======== Photo to Word by format ===========================
    print("\nGenerating...\n")
    recordsdir = []
    for root, dirs, files in os.walk(buffer_Path):
        for dir in dirs:
            recordsdir.append(dir)

    for dirName in recordsdir:
        document = Document()

        # page set up size
        section = document.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)

        # font set up style
        style = document.styles['Normal']
        font = style.font
        font.name = 'TH Sarabun New'
        font.size = Pt(16)
        font.bold = True

        # set up paragraph
        p = style.paragraph_format
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        p.line_spacing = 1

        # Insert Field Code to define data
        fieldcode = dirName
        records = []
        
        path_fieldcode = buffer_Path + fieldcode + "\\"

        # create table
        for root, dirs, files in os.walk(path_fieldcode):
            for filename in files:
                records.append(filename)
        data_size = int(len(records)/2)

        count_Table = 0
        p = document.add_paragraph('รูปถ่ายประกอบจุดรังวัด...............'+fieldcode+'...............')
        p.alignment = 1
        for i in range(data_size):
            document.add_paragraph('')
            table = document.add_table(rows=2, cols=2)
            table.style = "Table Grid"
            row0 = table.rows[0]
            row0.height = Mm(60)
            hdr_cells = row0.cells

            for j in range(2):
                paragraph = hdr_cells[j].paragraphs[0]
                run = paragraph.add_run()
                im = Image.open(path_fieldcode + records[j+i])
                width, height = im.size
                if (width > height):
                    run.add_picture(path_fieldcode + records[j+i],width = Mm(75),height = Mm(56.2))
                else:
                    run.add_picture(path_fieldcode + records[j+i],width = Mm(40),height = Mm(56.2))
            
            cell00 = table.cell(0,0)
            cell01 = table.cell(0,1)
            cell00.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell01.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell10 = table.cell(1,0)
            cell11 = table.cell(1,1)
            if(count_Table == 0):
                cell10.text = 'รูปจุดรังวัด'
                cell11.text = 'รูปเลขหมุดรังวัด'
                count_Table = count_Table + 1
            elif(count_Table == 1):
                cell10.text = 'รูปทิศเหนือ'
                cell11.text = 'รูปทิศใต้'
                count_Table = count_Table + 1
            elif(count_Table == 2):
                cell10.text = 'รูปทิศตะวันออก'
                cell11.text = 'รูปทิศตะวันตก'
                count_Table = 0
                if(i!=(data_size-1)):
                    document.add_page_break()
                    p = document.add_paragraph('รูปถ่ายประกอบจุดรังวัด...............'+fieldcode+'...............')
                    p.alignment = 1
            cell10.paragraphs[0].paragraph_format.alignment = 1
            cell11.paragraphs[0].paragraph_format.alignment = 1

            del(records[0])
        im.close()
        document.save(output_path + fieldcode + '.docx')
        print("\t" + output_path + fieldcode + '.docx' + " is completed.")
    print("\nDone. The job was completed.\n")

    shutil.rmtree(buffer_Path)
    input("Press Enter to continue...")

if __name__ == "__main__":
    main()